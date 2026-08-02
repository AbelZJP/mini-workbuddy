from __future__ import annotations

import shutil
import subprocess
import tempfile
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path
from typing import Any, Iterable


MAX_EXTRACTED_CHARS = 120_000
TEXT_SUFFIXES = {'.txt', '.md', '.json', '.csv'}
SUPPORTED_SUFFIXES = TEXT_SUFFIXES | {'.doc', '.docx', '.pdf', '.ppt', '.pptx', '.xlsx'}


class DocumentParseError(RuntimeError):
    """A user-facing error raised when a document cannot be extracted."""


def _clip(text: str) -> tuple[str, bool]:
    normalized = text.replace('\x00', '').strip()
    if len(normalized) <= MAX_EXTRACTED_CHARS:
        return normalized, False
    return normalized[:MAX_EXTRACTED_CHARS] + '\n\n[文档内容过长，后续内容已截断]', True


def _require_dependency(module_name: str, package_name: str) -> Any:
    try:
        return __import__(module_name)
    except ImportError as exc:
        raise DocumentParseError(
            f'当前环境缺少 {package_name}，请在项目虚拟环境中执行：'
            f'pip install {package_name}'
        ) from exc


def _parse_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise DocumentParseError('当前环境缺少 python-docx，请执行：pip install python-docx') from exc

    document = Document(str(path))
    sections: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            sections.append(text)
    for table_index, table in enumerate(document.tables, start=1):
        rows = []
        for row in table.rows:
            cells = [cell.text.replace('\n', ' ').strip() for cell in row.cells]
            if any(cells):
                rows.append(' | '.join(cells))
        if rows:
            sections.append(f'[表格 {table_index}]\n' + '\n'.join(rows))
    return '\n\n'.join(sections)


def _parse_text(path: Path) -> str:
    """读取可作为参考输入的纯文本文件。"""
    return path.read_text(encoding='utf-8', errors='replace')


def _parse_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise DocumentParseError('当前环境缺少 pypdf，请执行：pip install pypdf') from exc

    try:
        reader = PdfReader(str(path))
        if reader.is_encrypted:
            try:
                decrypted = reader.decrypt('')
            except Exception as exc:
                raise DocumentParseError('PDF 已加密，当前未提供密码，无法读取。') from exc
            if not decrypted:
                raise DocumentParseError('PDF 已加密，当前未提供密码，无法读取。')
        pages = []
        for index, page in enumerate(reader.pages, start=1):
            text = (page.extract_text() or '').strip()
            if text:
                pages.append(f'[第 {index} 页]\n{text}')
        return '\n\n'.join(pages)
    except DocumentParseError:
        raise
    except Exception as exc:
        raise DocumentParseError(f'PDF 解析失败：{exc}') from exc


def _cell_text(value: Any) -> str:
    if value is None:
        return ''
    if isinstance(value, float) and value.is_integer():
        return str(int(value))
    return str(value).replace('\n', ' ').strip()


def _parse_xlsx(path: Path) -> str:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise DocumentParseError('当前环境缺少 openpyxl，请执行：pip install openpyxl') from exc

    try:
        workbook = load_workbook(str(path), read_only=True, data_only=False)
        sheets: list[str] = []
        for worksheet in workbook.worksheets:
            rows: list[str] = []
            for values in worksheet.iter_rows(values_only=True):
                cells = [_cell_text(value) for value in values]
                while cells and not cells[-1]:
                    cells.pop()
                if any(cells):
                    rows.append(' | '.join(cells))
            if rows:
                sheets.append(f'[工作表：{worksheet.title}]\n' + '\n'.join(rows))
        workbook.close()
        return '\n\n'.join(sheets)
    except Exception as exc:
        raise DocumentParseError(f'Excel 文件解析失败：{exc}') from exc


def _parse_pptx(path: Path) -> str:
    """从 PPTX 的幻灯片 XML 提取文本，避免把 ZIP 二进制交给模型。"""
    try:
        with zipfile.ZipFile(path) as archive:
            slide_names = sorted(
                (
                    name for name in archive.namelist()
                    if name.startswith('ppt/slides/slide') and name.endswith('.xml')
                ),
                key=lambda name: int(Path(name).stem.removeprefix('slide'))
                if Path(name).stem.removeprefix('slide').isdigit() else name,
            )
            sections: list[str] = []
            text_tag = '{http://schemas.openxmlformats.org/drawingml/2006/main}t'
            for index, name in enumerate(slide_names, start=1):
                root = ET.fromstring(archive.read(name))
                text = ' '.join(
                    node.text.strip()
                    for node in root.iter(text_tag)
                    if node.text and node.text.strip()
                )
                if text:
                    sections.append(f'[第 {index} 页]\n{text}')
            return '\n\n'.join(sections)
    except zipfile.BadZipFile as exc:
        raise DocumentParseError('PPTX 文件结构损坏或不是有效的 PowerPoint 文件。') from exc
    except (KeyError, ET.ParseError, OSError) as exc:
        raise DocumentParseError(f'PPTX 解析失败：{exc}') from exc


def _parse_ppt(path: Path) -> str:
    """使用本机已有转换器提取旧版二进制 .ppt 的正文。"""
    textutil = shutil.which('textutil')
    if textutil:
        text = _run_text_converter([textutil, '-convert', 'txt', '-stdout', str(path)])
        if text:
            return text

    office = shutil.which('soffice') or shutil.which('libreoffice')
    if office:
        with tempfile.TemporaryDirectory(prefix='mini-workbuddy-ppt-') as temp_dir:
            output_dir = Path(temp_dir)
            try:
                completed = subprocess.run(
                    [office, '--headless', '--convert-to', 'txt:Text', '--outdir', str(output_dir), str(path)],
                    capture_output=True,
                    text=True,
                    timeout=60,
                    check=False,
                )
            except subprocess.TimeoutExpired:
                completed = None
            output_path = output_dir / f'{path.stem}.txt'
            if completed and completed.returncode == 0 and output_path.exists():
                text = output_path.read_text(encoding='utf-8', errors='ignore').strip()
                if text:
                    return text

    raise DocumentParseError(
        '无法解析 .ppt 文件：请安装 LibreOffice（soffice/libreoffice），或在 macOS 使用系统 textutil。'
    )


def _run_text_converter(command: list[str], timeout: int = 30) -> str:
    try:
        completed = subprocess.run(command, capture_output=True, text=True, timeout=timeout, check=False)
    except (FileNotFoundError, subprocess.TimeoutExpired):
        return ''
    if completed.returncode != 0:
        return ''
    return completed.stdout.strip()


def _parse_doc(path: Path) -> str:
    """Extract legacy binary .doc via an installed office conversion tool."""
    converters = []
    textutil = shutil.which('textutil')
    if textutil:
        converters.append([textutil, '-convert', 'txt', '-stdout', str(path)])
    antiword = shutil.which('antiword')
    if antiword:
        converters.append([antiword, str(path)])

    for command in converters:
        text = _run_text_converter(command)
        if text:
            return text

    office = shutil.which('soffice') or shutil.which('libreoffice')
    if office:
        with tempfile.TemporaryDirectory(prefix='mini-workbuddy-doc-') as temp_dir:
            output_dir = Path(temp_dir)
            try:
                completed = subprocess.run(
                    [office, '--headless', '--convert-to', 'txt:Text', '--outdir', str(output_dir), str(path)],
                    capture_output=True,
                    text=True,
                    timeout=45,
                    check=False,
                )
            except subprocess.TimeoutExpired:
                completed = None
            output_path = output_dir / f'{path.stem}.txt'
            if completed and completed.returncode == 0 and output_path.exists():
                text = output_path.read_text(encoding='utf-8', errors='ignore').strip()
                if text:
                    return text

    raise DocumentParseError(
        '无法解析 .doc 文件：请安装 LibreOffice（soffice）、antiword，或在 macOS 上使用系统 textutil。'
    )


def parse_document(path: Path) -> dict[str, Any]:
    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_SUFFIXES:
        raise DocumentParseError(f'暂不支持解析 {suffix or "该"} 文件。')
    if not path.exists() or not path.is_file():
        raise DocumentParseError('附件文件不存在，可能已被移动或删除。')

    if suffix in TEXT_SUFFIXES:
        text = _parse_text(path)
    elif suffix == '.docx':
        text = _parse_docx(path)
    elif suffix == '.doc':
        text = _parse_doc(path)
    elif suffix == '.pdf':
        text = _parse_pdf(path)
    elif suffix == '.pptx':
        text = _parse_pptx(path)
    elif suffix == '.ppt':
        text = _parse_ppt(path)
    else:
        text = _parse_xlsx(path)
    clipped_text, truncated = _clip(text)
    return {'format': suffix[1:].upper(), 'text': clipped_text, 'truncated': truncated}


def parse_attachments(workspace_root: Path, attachments: Iterable[dict[str, Any]]) -> tuple[str, list[dict[str, Any]]]:
    """Parse supported attachments while enforcing the current workspace boundary."""
    root = workspace_root.expanduser().resolve()
    blocks: list[str] = []
    metadata: list[dict[str, Any]] = []
    for attachment in attachments:
        raw_path = str(attachment.get('path') or '').strip()
        if not raw_path:
            continue
        candidate = (root / raw_path).resolve()
        relative = candidate.relative_to(root).as_posix() if candidate.is_relative_to(root) else raw_path
        item = {'path': relative, 'name': str(attachment.get('name') or Path(relative).name), 'status': 'skipped'}
        if not candidate.is_relative_to(root):
            item.update({'status': 'failed', 'error': '附件路径超出当前工作空间，已拒绝读取。'})
            metadata.append(item)
            blocks.append(f'附件 {relative} 读取失败：附件路径超出当前工作空间，已拒绝读取。')
            continue
        if candidate.suffix.lower() not in SUPPORTED_SUFFIXES:
            metadata.append(item)
            continue
        try:
            parsed = parse_document(candidate)
            item.update({'status': 'parsed', 'format': parsed['format'], 'truncated': parsed['truncated']})
            metadata.append(item)
            content = parsed['text'] or '（文件中没有提取到可用文本，可能是扫描件或图片型文档。）'
            blocks.append(f'附件文件：{relative}\n文件格式：{parsed["format"]}\n文件内容：\n{content}')
        except DocumentParseError as exc:
            item.update({'status': 'failed', 'error': str(exc)})
            metadata.append(item)
            blocks.append(f'附件 {relative} 解析失败：{exc}')
    return '\n\n---\n\n'.join(blocks), metadata
