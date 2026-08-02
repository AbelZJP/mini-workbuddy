from pathlib import Path
from tempfile import TemporaryDirectory
import zipfile
import unittest

from docx import Document
from openpyxl import Workbook

from .document_parser import parse_attachments, parse_document


def _write_pdf(path: Path) -> None:
    objects = [
        b'<< /Type /Catalog /Pages 2 0 R >>',
        b'<< /Type /Pages /Kids [3 0 R] /Count 1 >>',
        b'<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Resources << /Font << /F1 5 0 R >> >> /Contents 4 0 R >>',
        b'<< /Length 47 >>\nstream\nBT /F1 12 Tf 72 720 Td (PDF parser test) Tj ET\nendstream',
        b'<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>',
    ]
    output = bytearray(b'%PDF-1.4\n')
    offsets = [0]
    for index, obj in enumerate(objects, 1):
        offsets.append(len(output))
        output.extend(f'{index} 0 obj\n'.encode())
        output.extend(obj)
        output.extend(b'\nendobj\n')
    xref = len(output)
    output.extend(f'xref\n0 {len(objects) + 1}\n'.encode())
    output.extend(b'0000000000 65535 f \n')
    for offset in offsets[1:]:
        output.extend(f'{offset:010d} 00000 n \n'.encode())
    output.extend(f'trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref}\n%%EOF\n'.encode())
    path.write_bytes(output)


def _write_pptx(path: Path) -> None:
    slide = '''
    <p:sld xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main"
           xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">
      <p:cSld><p:spTree><p:sp><p:txBody><a:p><a:r><a:t>季度汇报</a:t></a:r></a:p>
      <a:p><a:r><a:t>本季度完成情况</a:t></a:r></a:p></p:txBody></p:sp></p:spTree></p:cSld>
    </p:sld>
    '''
    with zipfile.ZipFile(path, 'w') as archive:
        archive.writestr('ppt/slides/slide1.xml', slide)


class DocumentParserTest(unittest.TestCase):
    def test_docx_pdf_pptx_xlsx_and_workspace_boundary(self) -> None:
        with TemporaryDirectory() as temp:
            root = Path(temp)

            text_files = {
                '.txt': '项目目标：生成一份包含市场背景和实施计划的演示文稿。',
                '.md': '# 参考资料\n\n需要保留实施计划。',
                '.json': '{"主题": "市场背景"}',
                '.csv': '章节,内容\n市场背景,行业规模',
            }
            for suffix, content in text_files.items():
                (root / f'参考资料{suffix}').write_text(content, encoding='utf-8')

            docx_path = root / '合同.docx'
            document = Document()
            document.add_paragraph('甲方应按约付款。')
            table = document.add_table(rows=1, cols=2)
            table.rows[0].cells[0].text = '期限'
            table.rows[0].cells[1].text = '30天'
            document.save(docx_path)

            xlsx_path = root / '数据.xlsx'
            workbook = Workbook()
            worksheet = workbook.active
            worksheet.title = '销售'
            worksheet.append(['月份', '金额'])
            worksheet.append(['1月', 100])
            workbook.save(xlsx_path)

            pdf_path = root / '说明.pdf'
            _write_pdf(pdf_path)

            pptx_path = root / '季度汇报.pptx'
            _write_pptx(pptx_path)

            self.assertIn('甲方应按约付款。', parse_document(docx_path)['text'])
            self.assertIn('销售', parse_document(xlsx_path)['text'])
            self.assertIn('PDF parser test', parse_document(pdf_path)['text'])
            self.assertEqual(parse_document(pptx_path)['format'], 'PPTX')
            self.assertIn('季度汇报', parse_document(pptx_path)['text'])
            for suffix in text_files:
                parsed_text = parse_document(root / f'参考资料{suffix}')
                self.assertEqual(parsed_text['format'], suffix[1:].upper())
                self.assertTrue(parsed_text['text'])
            self.assertIn('市场背景', parse_document(root / '参考资料.txt')['text'])

            text, metadata = parse_attachments(root, [
                {'path': '合同.docx', 'name': '合同.docx'},
                {'path': '季度汇报.pptx', 'name': '季度汇报.pptx'},
                {'path': '../outside.txt', 'name': 'outside.txt'},
            ])
            self.assertIn('甲方应按约付款。', text)
            self.assertEqual(metadata[0]['status'], 'parsed')
            self.assertEqual(metadata[1]['status'], 'parsed')
            self.assertEqual(metadata[2]['status'], 'failed')


if __name__ == '__main__':
    unittest.main()
